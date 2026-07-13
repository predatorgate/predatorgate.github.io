# -*- coding: utf-8 -*-
"""Build a .docx with a 4-column table from gr/timeline.csv + sources.json."""

import csv
import hashlib
import json
import re

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, Cm


def make_entry_id(date, description):
    raw = f"{date}|{description}"
    return hashlib.sha256(raw.encode("utf-8")).hexdigest()[:8]


def to_iso_date(d):
    m = re.fullmatch(r"(\d{1,2})/(\d{1,2})/(\d{4})", d.strip())
    if not m:
        return d
    day, month, year = m.groups()
    return f"{year}-{int(month):02d}-{int(day):02d}"


def add_hyperlink(paragraph, url, text):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    rStyle = OxmlElement("w:rStyle")
    rStyle.set(qn("w:val"), "Hyperlink")
    rPr.append(rStyle)

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0000FF")
    rPr.append(color)

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rPr.append(underline)

    new_run.append(rPr)

    t = OxmlElement("w:t")
    t.text = text
    t.set(qn("xml:space"), "preserve")
    new_run.append(t)

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink


def main():
    with open("sources.json") as f:
        sources = json.load(f)

    with open("gr/timeline.csv", newline="") as f:
        rows = list(csv.reader(f))

    data_rows = []
    for row in rows[1:]:
        if len(row) < 3:
            continue
        date, event, _url = row[0], row[1], row[2]
        if not date and not event:
            continue
        key = make_entry_id(date, event)
        urls = sources.get(key, [])
        data_rows.append((date, event, urls))

    doc = Document()

    # Landscape A4 with narrow margins to give the table room.
    section = doc.sections[0]
    section.page_height, section.page_width = section.page_width, section.page_height
    section.left_margin = Cm(1.2)
    section.right_margin = Cm(1.2)
    section.top_margin = Cm(1.2)
    section.bottom_margin = Cm(1.2)

    headings = ["Ημερομηνία", "Γεγονός", "Πηγές", "Λέξεις-κλειδιά"]
    table = doc.add_table(rows=1 + len(data_rows), cols=4)
    table.style = "Light Grid"
    table.autofit = False

    widths = [Cm(2.5), Cm(12.0), Cm(9.0), Cm(4.0)]
    for i, w in enumerate(widths):
        for cell in table.columns[i].cells:
            cell.width = w

    hdr = table.rows[0].cells
    for i, h in enumerate(headings):
        p = hdr[i].paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(11)

    for r_idx, (date, event, urls) in enumerate(data_rows, start=1):
        cells = table.rows[r_idx].cells
        cells[0].text = ""
        cells[0].paragraphs[0].add_run(to_iso_date(date)).font.size = Pt(10)
        cells[1].text = ""
        cells[1].paragraphs[0].add_run(event).font.size = Pt(10)

        src_cell = cells[2]
        src_cell.text = ""
        first = True
        for url in urls:
            p = src_cell.paragraphs[0] if first else src_cell.add_paragraph()
            first = False
            p.style = doc.styles["List Bullet"]
            add_hyperlink(p, url, url)
            for run in p.runs:
                run.font.size = Pt(9)

        cells[3].text = ""

        if r_idx % 50 == 0:
            print(f"  wrote {r_idx}/{len(data_rows)} rows")

    out = "timeline_gr.docx"
    doc.save(out)
    print(f"Saved {out} with {len(data_rows)} entries")


if __name__ == "__main__":
    main()
